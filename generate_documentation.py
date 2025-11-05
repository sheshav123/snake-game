from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_heading_custom(doc, text, level=1, color=RGBColor(0, 0, 0)):
    """Add a custom styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
    return heading

def add_centered_text(doc, text, bold=False, size=14, color=RGBColor(0, 0, 0)):
    """Add centered text"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return para

def add_horizontal_line(doc):
    """Add a horizontal line"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('_' * 80)
    run.font.color.rgb = RGBColor(128, 128, 128)

def create_documentation():
    """Create comprehensive Snake Game documentation"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ==================== CERTIFICATE PAGE ====================
    doc.add_paragraph('\n' * 8)
    add_centered_text(doc, 'CERTIFICATE', bold=True, size=24, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n' * 2)
    add_horizontal_line(doc)
    doc.add_paragraph()
    
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cert_text.add_run(
        'This is to certify that the project entitled '
    )
    cert_text.add_run('"SNAKE GAME USING PYTHON AND PYGAME"').bold = True
    cert_text.add_run(
        ' submitted in partial fulfillment of the requirements for the degree of '
    )
    cert_text.add_run('Master of Computer Applications (MCA)').bold = True
    cert_text.add_run(' is a bonafide work carried out during the academic year ')
    cert_text.add_run('2025').bold = True
    cert_text.add_run('.')
    
    doc.add_paragraph('\n' * 3)
    
    # Signature section
    sig_para = doc.add_paragraph()
    sig_para.add_run('\n\n\n')
    
    left_sig = doc.add_paragraph()
    left_sig.add_run('Project Guide').bold = True
    left_sig.add_run('\n_____________________')
    
    doc.add_paragraph('\n' * 2)
    
    right_sig = doc.add_paragraph()
    right_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_sig.add_run('Head of Department').bold = True
    right_sig.add_run('\n_____________________')
    
    doc.add_paragraph('\n' * 2)
    date_para = doc.add_paragraph()
    date_para.add_run(f'Date: {datetime.datetime.now().strftime("%B %d, %Y")}')
    
    add_page_break(doc)
    
    # ==================== ACKNOWLEDGEMENT ====================
    doc.add_paragraph('\n' * 2)
    add_heading_custom(doc, 'ACKNOWLEDGEMENT', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    ack_para1 = doc.add_paragraph()
    ack_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ack_para1.add_run(
        'I would like to express my sincere gratitude to all those who have contributed to the '
        'successful completion of this project. First and foremost, I am deeply thankful to my '
        'project guide for their invaluable guidance, constant encouragement, and constructive '
        'feedback throughout the development of this Snake Game application.'
    )
    
    doc.add_paragraph()
    
    ack_para2 = doc.add_paragraph()
    ack_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ack_para2.add_run(
        'I extend my heartfelt thanks to the Head of the Department and all faculty members of '
        'the Computer Applications department for providing the necessary resources and creating '
        'an environment conducive to learning and innovation.'
    )
    
    doc.add_paragraph()
    
    ack_para3 = doc.add_paragraph()
    ack_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ack_para3.add_run(
        'I am also grateful to my classmates and peers for their support, suggestions, and '
        'collaborative spirit during the course of this project. Their insights have been '
        'instrumental in refining various aspects of the game.'
    )
    
    doc.add_paragraph()
    
    ack_para4 = doc.add_paragraph()
    ack_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ack_para4.add_run(
        'Finally, I would like to thank my family for their unwavering support and encouragement, '
        'which has been a constant source of motivation throughout my academic journey.'
    )
    
    doc.add_paragraph('\n' * 3)
    
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run('Student Name').bold = True
    signature.add_run('\nMCA Semester 3')
    
    add_page_break(doc)
    
    # ==================== TABLE OF CONTENTS ====================
    doc.add_paragraph('\n' * 2)
    add_heading_custom(doc, 'TABLE OF CONTENTS', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    toc_items = [
        ('1.', 'Certificate', '1'),
        ('2.', 'Acknowledgement', '2'),
        ('3.', 'Abstract', '4'),
        ('4.', 'Introduction', '5'),
        ('5.', 'Objectives', '6'),
        ('6.', 'About the Project', '7'),
        ('7.', 'System Analysis', '8'),
        ('   7.1', 'Hardware Requirements', '8'),
        ('   7.2', 'Software Requirements', '8'),
        ('   7.3', 'Functional Requirements', '9'),
        ('8.', 'Flowchart', '10'),
        ('9.', 'Code Implementation', '11'),
        ('10.', 'Output Screenshots', '15'),
        ('11.', 'Real-World Use Cases', '16'),
        ('12.', 'Future Enhancements', '17'),
        ('13.', 'Future Scope', '18'),
        ('14.', 'Summary', '19'),
    ]
    
    for num, title, page in toc_items:
        toc_para = doc.add_paragraph()
        toc_para.add_run(f'{num} {title}').bold = True
        toc_para.add_run('.' * (70 - len(num) - len(title) - len(page)))
        toc_para.add_run(f' {page}')
    
    add_page_break(doc)
    
    # ==================== ABSTRACT ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, 'ABSTRACT', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    abstract_para1 = doc.add_paragraph()
    abstract_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_para1.add_run(
        'This project presents the development of a classic Snake Game using Python programming '
        'language and the Pygame library. The Snake Game is a timeless arcade game that has been '
        'popular since its inception in the late 1970s and gained massive popularity on mobile '
        'devices in the late 1990s.'
    )
    
    doc.add_paragraph()
    
    abstract_para2 = doc.add_paragraph()
    abstract_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_para2.add_run(
        'The implementation utilizes object-oriented programming principles to create a modular, '
        'maintainable, and extensible codebase. The game features smooth gameplay mechanics, '
        'collision detection, score tracking, pause functionality, and an intuitive user interface. '
        'The snake navigates through a grid-based playing field, consuming food items to grow longer '
        'while avoiding collisions with its own body.'
    )
    
    doc.add_paragraph()
    
    abstract_para3 = doc.add_paragraph()
    abstract_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_para3.add_run(
        'The project demonstrates fundamental concepts of game development including game loops, '
        'event handling, sprite rendering, state management, and real-time user input processing. '
        'Pygame library provides the necessary tools for graphics rendering, event handling, and '
        'game timing, making it an ideal choice for 2D game development in Python.'
    )
    
    doc.add_paragraph()
    
    abstract_para4 = doc.add_paragraph()
    abstract_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_para4.add_run(
        'This documentation provides comprehensive coverage of the project including system requirements, '
        'design methodology, code implementation, testing procedures, and future enhancement possibilities. '
        'The game serves as an excellent educational tool for understanding game development concepts and '
        'Python programming practices.'
    )
    
    add_page_break(doc)
    
    # ==================== INTRODUCTION ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '1. INTRODUCTION', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    intro_para1 = doc.add_paragraph()
    intro_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_para1.add_run(
        'The Snake Game is one of the most iconic and enduring video games in history. Originally '
        'created in the 1970s, it gained worldwide recognition when Nokia featured it on their mobile '
        'phones in 1998. The game\'s simple yet addictive gameplay has made it a perfect candidate for '
        'learning game development concepts.'
    )
    
    doc.add_paragraph()
    
    intro_para2 = doc.add_paragraph()
    intro_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_para2.add_run(
        'In this project, we have developed a modern implementation of the classic Snake Game using '
        'Python 3 and the Pygame library. Python is an excellent choice for game development due to '
        'its simplicity, readability, and extensive library support. Pygame, built on top of the SDL '
        '(Simple DirectMedia Layer) library, provides powerful tools for handling graphics, sound, '
        'and user input.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('1.1 Game Concept', level=2)
    concept_para = doc.add_paragraph()
    concept_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    concept_para.add_run(
        'The game involves controlling a snake that moves continuously in a grid-based environment. '
        'The player controls the direction of the snake using arrow keys. Food items appear randomly '
        'on the grid, and when the snake consumes food, it grows longer and the player\'s score increases. '
        'The game ends when the snake collides with itself. The objective is to achieve the highest '
        'possible score by consuming as much food as possible without self-collision.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('1.2 Motivation', level=2)
    motivation_para = doc.add_paragraph()
    motivation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    motivation_para.add_run(
        'The motivation behind this project is multifaceted: (1) To understand and implement fundamental '
        'game development concepts, (2) To practice object-oriented programming in Python, (3) To gain '
        'hands-on experience with the Pygame library, (4) To develop problem-solving skills through '
        'algorithm implementation, and (5) To create an entertaining and educational application that '
        'demonstrates programming proficiency.'
    )
    
    add_page_break(doc)
    
    # ==================== OBJECTIVES ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '2. OBJECTIVES', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    obj_intro = doc.add_paragraph()
    obj_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obj_intro.add_run('The primary objectives of this project are:')
    
    doc.add_paragraph()
    
    objectives = [
        ('Primary Objectives:', [
            'To develop a fully functional Snake Game with smooth gameplay mechanics',
            'To implement collision detection algorithms for game-over conditions',
            'To create an intuitive user interface with real-time score display',
            'To utilize object-oriented programming principles for code organization',
        ]),
        ('Technical Objectives:', [
            'To implement a game loop for continuous gameplay',
            'To handle keyboard events for real-time user input',
            'To render graphics efficiently using Pygame',
            'To manage game states (playing, paused, game over)',
            'To implement random food generation algorithm',
        ]),
        ('Learning Objectives:', [
            'To understand game development lifecycle and architecture',
            'To practice Python programming with type hints and enumerations',
            'To learn event-driven programming concepts',
            'To implement data structures (lists, tuples) for game logic',
            'To develop debugging and testing skills',
        ]),
        ('User Experience Objectives:', [
            'To provide responsive and intuitive controls',
            'To create visually appealing graphics with proper color schemes',
            'To implement pause and restart functionality',
            'To display clear game status messages',
        ]),
    ]
    
    for category, items in objectives:
        doc.add_heading(category, level=3)
        for item in items:
            para = doc.add_paragraph(item, style='List Bullet')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # ==================== ABOUT THE PROJECT ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '3. ABOUT THE PROJECT', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    doc.add_heading('3.1 Project Overview', level=2)
    overview_para = doc.add_paragraph()
    overview_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    overview_para.add_run(
        'This Snake Game project is a desktop application built using Python 3 and Pygame. The game '
        'features a 800x600 pixel window with a grid-based playing field divided into 20x20 pixel cells. '
        'The snake starts with an initial length of 2 segments and grows by one segment each time it '
        'consumes food. The game runs at 10 frames per second, providing smooth and controllable gameplay.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Key Features', level=2)
    features = [
        'Grid-based Movement: The snake moves in discrete grid cells for precise control',
        'Collision Detection: Accurate detection of snake-body collisions',
        'Score Tracking: Real-time score display showing the number of food items consumed',
        'Pause Functionality: Ability to pause and resume the game using the P key',
        'Restart Option: Quick restart after game over using the R key',
        'Wrap-around Borders: Snake wraps around screen edges instead of hitting walls',
        'Visual Feedback: Different colors for snake head, body, and food',
        'Grid Display: Subtle grid lines for better spatial awareness',
        'Responsive Controls: Immediate response to arrow key inputs',
        'Game State Management: Proper handling of different game states',
    ]
    
    for feature in features:
        para = doc.add_paragraph(feature, style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Technology Stack', level=2)
    
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_data = [
        ('Programming Language', 'Python 3.x'),
        ('Game Library', 'Pygame 2.x'),
        ('Development Paradigm', 'Object-Oriented Programming'),
        ('Type System', 'Python Type Hints'),
        ('Standard Libraries', 'random, time, sys, enum, typing'),
    ]
    
    for i, (key, value) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = key
        tech_table.rows[i].cells[1].text = value
        tech_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('3.4 Game Mechanics', level=2)
    mechanics_para = doc.add_paragraph()
    mechanics_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mechanics_para.add_run(
        'The game implements several core mechanics: (1) Continuous Movement - the snake moves '
        'automatically in the current direction, (2) Direction Control - players can change direction '
        'but cannot reverse directly, (3) Growth Mechanism - consuming food adds one segment to the tail, '
        '(4) Collision System - the game ends when the snake\'s head touches any part of its body, '
        '(5) Food Spawning - food appears randomly at positions not occupied by the snake.'
    )
    
    add_page_break(doc)
    
    # ==================== SYSTEM ANALYSIS ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '4. SYSTEM ANALYSIS', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    doc.add_heading('4.1 Hardware Requirements', level=2)
    
    hw_table = doc.add_table(rows=5, cols=2)
    hw_table.style = 'Light Grid Accent 1'
    
    hw_data = [
        ('Processor', 'Intel Core i3 or equivalent (minimum)\nIntel Core i5 or higher (recommended)'),
        ('RAM', '2 GB (minimum)\n4 GB or higher (recommended)'),
        ('Storage', '100 MB free disk space'),
        ('Display', '800x600 resolution or higher'),
        ('Input Device', 'Keyboard with arrow keys'),
    ]
    
    for i, (key, value) in enumerate(hw_data):
        hw_table.rows[i].cells[0].text = key
        hw_table.rows[i].cells[1].text = value
        hw_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Software Requirements', level=2)
    
    sw_table = doc.add_table(rows=5, cols=2)
    sw_table.style = 'Light Grid Accent 1'
    
    sw_data = [
        ('Operating System', 'Windows 10/11, macOS 10.12+, or Linux (Ubuntu 18.04+)'),
        ('Python Version', 'Python 3.7 or higher'),
        ('Pygame Library', 'Pygame 2.0.0 or higher'),
        ('Development IDE', 'Any Python IDE (VS Code, PyCharm, etc.) - Optional'),
        ('Additional Libraries', 'Standard Python libraries (random, time, sys, enum, typing)'),
    ]
    
    for i, (key, value) in enumerate(sw_data):
        sw_table.rows[i].cells[0].text = key
        sw_table.rows[i].cells[1].text = value
        sw_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 Functional Requirements', level=2)
    
    func_reqs = [
        'The system shall allow the player to control snake direction using arrow keys',
        'The system shall detect collisions between the snake head and its body',
        'The system shall generate food at random positions not occupied by the snake',
        'The system shall increase the snake length by one segment when food is consumed',
        'The system shall increment the score by one point for each food item consumed',
        'The system shall display the current score in real-time',
        'The system shall allow the player to pause and resume the game',
        'The system shall allow the player to restart the game after game over',
        'The system shall prevent the snake from reversing direction directly',
        'The system shall wrap the snake around screen edges',
        'The system shall display appropriate messages for game states',
        'The system shall maintain a consistent frame rate for smooth gameplay',
    ]
    
    for req in func_reqs:
        para = doc.add_paragraph(req, style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    doc.add_heading('4.4 Non-Functional Requirements', level=2)
    
    non_func_reqs = [
        'Performance: The game shall run at a consistent 10 FPS without lag',
        'Usability: The game controls shall be intuitive and responsive',
        'Reliability: The game shall not crash during normal gameplay',
        'Maintainability: The code shall be well-organized and documented',
        'Portability: The game shall run on Windows, macOS, and Linux',
        'Scalability: The code structure shall allow easy addition of new features',
    ]
    
    for req in non_func_reqs:
        para = doc.add_paragraph(req, style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    add_page_break(doc)
    
    # ==================== FLOWCHART ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '5. FLOWCHART', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    doc.add_heading('5.1 Main Game Loop Flowchart', level=2)
    
    flowchart_desc = doc.add_paragraph()
    flowchart_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    flowchart_desc.add_run(
        'The following flowchart illustrates the main game loop logic:'
    )
    
    doc.add_paragraph()
    
    # Flowchart steps in text format
    flow_steps = [
        'START',
        '↓',
        'Initialize Pygame and Game Window',
        '↓',
        'Create SnakeGame Instance',
        '↓',
        'Reset Game State (snake position, direction, score)',
        '↓',
        'Generate Initial Food Position',
        '↓',
        'GAME LOOP START',
        '↓',
        'Handle Events (keyboard input, quit)',
        '↓',
        'Is Quit Event? → YES → END GAME',
        '↓ NO',
        'Is Game Paused? → YES → Skip Update',
        '↓ NO',
        'Is Game Over? → YES → Wait for Restart',
        '↓ NO',
        'Update Snake Direction',
        '↓',
        'Calculate New Head Position',
        '↓',
        'Check Self-Collision → YES → Set Game Over',
        '↓ NO',
        'Add New Head to Snake',
        '↓',
        'Check Food Collision → YES → Increase Score, Generate New Food',
        '↓ NO',
        'Remove Tail Segment',
        '↓',
        'Draw Background and Grid',
        '↓',
        'Draw Snake',
        '↓',
        'Draw Food',
        '↓',
        'Draw Score',
        '↓',
        'Draw Game Status Messages',
        '↓',
        'Update Display',
        '↓',
        'Control Frame Rate (10 FPS)',
        '↓',
        'LOOP BACK TO GAME LOOP START',
    ]
    
    for step in flow_steps:
        para = doc.add_paragraph(step)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if step in ['START', 'END GAME', 'GAME LOOP START']:
            para.runs[0].font.bold = True
            para.runs[0].font.color.rgb = RGBColor(139, 0, 0)
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Collision Detection Logic', level=2)
    
    collision_desc = doc.add_paragraph()
    collision_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    collision_desc.add_run(
        'The collision detection follows this logic: After calculating the new head position, '
        'the system checks if this position exists in the current snake body list. If found, '
        'it indicates self-collision and triggers game over. For food collision, the system '
        'compares the new head position with the food position for exact match.'
    )
    
    add_page_break(doc)
    
    # ==================== CODE IMPLEMENTATION ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '6. CODE IMPLEMENTATION', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    doc.add_heading('6.1 Code Structure', level=2)
    
    structure_para = doc.add_paragraph()
    structure_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    structure_para.add_run(
        'The code is organized using object-oriented programming principles with a main SnakeGame '
        'class that encapsulates all game logic. The implementation uses Python enumerations for '
        'direction management and type hints for better code clarity.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Key Components', level=2)
    
    components = [
        ('Constants Section', 'Defines window dimensions, grid size, colors, and frame rate'),
        ('Direction Enum', 'Enumeration class for four movement directions with coordinate deltas'),
        ('SnakeGame Class', 'Main game class containing all game logic and rendering'),
        ('__init__ Method', 'Initializes Pygame, creates window, and sets up game components'),
        ('reset_game Method', 'Resets all game variables to initial state'),
        ('generate_food Method', 'Randomly generates food position avoiding snake body'),
        ('handle_events Method', 'Processes keyboard input and window events'),
        ('update Method', 'Updates game state including snake movement and collision detection'),
        ('draw Method', 'Renders all visual elements to the screen'),
        ('run Method', 'Main game loop that coordinates all game operations'),
    ]
    
    for component, description in components:
        para = doc.add_paragraph()
        para.add_run(f'{component}: ').bold = True
        para.add_run(description)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    doc.add_heading('6.3 Complete Source Code', level=2)
    
    code_intro = doc.add_paragraph()
    code_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    code_intro.add_run(
        'The complete source code is presented below with detailed inline comments explaining '
        'each section:'
    )
    
    doc.add_paragraph()
    
    # Add code with proper formatting
    code = '''import pygame
import random
import time
import sys
from enum import Enum
from typing import List, Tuple, Optional

# Initialize pygame
pygame.init()

# Constants - Define game window and grid parameters
WINDOW_WIDTH = 800
WINDOW_HEIGHT = 600
GRID_SIZE = 20
GRID_WIDTH = WINDOW_WIDTH // GRID_SIZE
GRID_HEIGHT = WINDOW_HEIGHT // GRID_SIZE
FPS = 10

# Colors - RGB color definitions
BLACK = (0, 0, 0)
WHITE = (255, 255, 255)
GREEN = (0, 255, 0)
RED = (255, 0, 0)
BLUE = (0, 0, 255)
DARK_GREEN = (0, 200, 0)
GRAY = (200, 200, 200)

# Directions - Enumeration for movement directions
class Direction(Enum):    
    UP = (0, -1)
    DOWN = (0, 1)
    LEFT = (-1, 0)
    RIGHT = (1, 0)

class SnakeGame:
    """Main game class containing all game logic"""
    
    def __init__(self):
        """Initialize game window and components"""
        self.screen = pygame.display.set_mode((WINDOW_WIDTH, WINDOW_HEIGHT))
        pygame.display.set_caption('Snake Game')
        self.clock = pygame.time.Clock()
        self.font = pygame.font.SysFont('Arial', 30)
        self.reset_game()

    def reset_game(self):
        """Reset game to initial state"""
        self.direction = Direction.RIGHT
        self.next_direction = Direction.RIGHT
        self.snake = [(GRID_WIDTH // 2, GRID_HEIGHT // 2)]
        self.snake.append((self.snake[0][0] - 1, self.snake[0][1]))
        self.food = self.generate_food()
        self.score = 0
        self.game_over = False
        self.paused = False

    def generate_food(self) -> Tuple[int, int]:
        """Generate random food position not on snake"""
        while True:
            food = (random.randint(0, GRID_WIDTH - 1), 
                   random.randint(0, GRID_HEIGHT - 1))
            if food not in self.snake:
                return food

    def handle_events(self):
        """Process keyboard and window events"""
        for event in pygame.event.get():
            if event.type == pygame.QUIT:
                return False
            
            if event.type == pygame.KEYDOWN:
                if event.key == pygame.K_ESCAPE:
                    return False
                elif event.key == pygame.K_p:
                    self.paused = not self.paused
                elif event.key == pygame.K_r and self.game_over:
                    self.reset_game()
                elif not self.game_over and not self.paused:
                    if event.key == pygame.K_UP and self.direction != Direction.DOWN:
                        self.next_direction = Direction.UP
                    elif event.key == pygame.K_DOWN and self.direction != Direction.UP:
                        self.next_direction = Direction.DOWN
                    elif event.key == pygame.K_LEFT and self.direction != Direction.RIGHT:
                        self.next_direction = Direction.LEFT
                    elif event.key == pygame.K_RIGHT and self.direction != Direction.LEFT:
                        self.next_direction = Direction.RIGHT
        
        return True

    def update(self):
        """Update game state"""
        if self.game_over or self.paused:
            return

        # Update direction
        self.direction = self.next_direction
        dx, dy = self.direction.value
        
        # Move snake
        head_x, head_y = self.snake[0]
        new_head = ((head_x + dx) % GRID_WIDTH, (head_y + dy) % GRID_HEIGHT)
        
        # Check for collision with self
        if new_head in self.snake:
            self.game_over = True
            return
        
        self.snake.insert(0, new_head)
        
        # Check if food is eaten
        if new_head == self.food:
            self.score += 1
            self.food = self.generate_food()
        else:
            self.snake.pop()

    def draw(self):
        """Render all game elements"""
        self.screen.fill(BLACK)
        
        # Draw grid
        for x in range(0, WINDOW_WIDTH, GRID_SIZE):
            pygame.draw.line(self.screen, (50, 50, 50), 
                           (x, 0), (x, WINDOW_HEIGHT))
        for y in range(0, WINDOW_HEIGHT, GRID_SIZE):
            pygame.draw.line(self.screen, (50, 50, 50), 
                           (0, y), (WINDOW_WIDTH, y))
        
        # Draw snake
        for i, (x, y) in enumerate(self.snake):
            color = DARK_GREEN if i == 0 else GREEN
            pygame.draw.rect(self.screen, color, 
                           (x * GRID_SIZE, y * GRID_SIZE, GRID_SIZE, GRID_SIZE))
            pygame.draw.rect(self.screen, (0, 100, 0), 
                           (x * GRID_SIZE, y * GRID_SIZE, GRID_SIZE, GRID_SIZE), 1)
        
        # Draw food
        pygame.draw.rect(self.screen, RED, 
                        (self.food[0] * GRID_SIZE, self.food[1] * GRID_SIZE, 
                         GRID_SIZE, GRID_SIZE))
        
        # Draw score
        score_text = self.font.render(f'Score: {self.score}', True, WHITE)
        self.screen.blit(score_text, (10, 10))
        
        # Draw game over or paused message
        if self.game_over:
            self.draw_centered_text('GAME OVER! Press R to restart', RED)
        elif self.paused:
            self.draw_centered_text('PAUSED - Press P to continue', WHITE)
        
        pygame.display.flip()
    
    def draw_centered_text(self, text, color):
        """Draw centered text on screen"""
        text_surface = self.font.render(text, True, color)
        text_rect = text_surface.get_rect(
            center=(WINDOW_WIDTH // 2, WINDOW_HEIGHT // 2))
        self.screen.blit(text_surface, text_rect)

    def run(self):
        """Main game loop"""
        running = True
        while running:
            running = self.handle_events()
            self.update()
            self.draw()
            self.clock.tick(FPS)
        
        pygame.quit()
        sys.exit()

if __name__ == "__main__":
    game = SnakeGame()
    game.run()
'''
    
    # Add code in monospace font
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    add_page_break(doc)
    
    # ==================== OUTPUT ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '7. OUTPUT SCREENSHOTS', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    doc.add_heading('7.1 Game States', level=2)
    
    output_desc = doc.add_paragraph()
    output_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    output_desc.add_run(
        'The game displays different visual states during gameplay:'
    )
    
    doc.add_paragraph()
    
    states = [
        ('Initial State', 'Game starts with a 2-segment snake in the center, moving right. '
         'Food appears at a random position. Score displays as 0.'),
        ('Playing State', 'Snake moves continuously in the current direction. Grid lines are '
         'visible. Snake head is dark green, body is bright green. Food is red. Current score '
         'is displayed in the top-left corner.'),
        ('Paused State', 'Game freezes with "PAUSED - Press P to continue" message displayed '
         'in white at the center of the screen. All game elements remain visible.'),
        ('Game Over State', 'When snake collides with itself, "GAME OVER! Press R to restart" '
         'message appears in red at the center. Final score remains displayed.'),
    ]
    
    for state, description in states:
        para = doc.add_paragraph()
        para.add_run(f'{state}: ').bold = True
        para.add_run(description)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
    
    doc.add_heading('7.2 Visual Elements', level=2)
    
    visual_table = doc.add_table(rows=6, cols=2)
    visual_table.style = 'Light Grid Accent 1'
    
    visual_data = [
        ('Background', 'Black (RGB: 0, 0, 0)'),
        ('Grid Lines', 'Dark Gray (RGB: 50, 50, 50)'),
        ('Snake Head', 'Dark Green (RGB: 0, 200, 0)'),
        ('Snake Body', 'Bright Green (RGB: 0, 255, 0)'),
        ('Food', 'Red (RGB: 255, 0, 0)'),
        ('Text/Score', 'White (RGB: 255, 255, 255)'),
    ]
    
    for i, (element, color) in enumerate(visual_data):
        visual_table.rows[i].cells[0].text = element
        visual_table.rows[i].cells[1].text = color
        visual_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('7.3 Sample Gameplay Scenario', level=2)
    
    scenario_para = doc.add_paragraph()
    scenario_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    scenario_para.add_run(
        'A typical gameplay session: The player starts the game and sees the snake in the center. '
        'Using arrow keys, they navigate toward the red food. Upon eating it, the score increases '
        'to 1 and the snake grows by one segment. A new food item appears elsewhere. The player '
        'continues eating food, and the snake grows longer with each consumption. Eventually, as '
        'the snake becomes very long, the player accidentally steers the head into the body, '
        'triggering the game over screen showing the final score.'
    )
    
    add_page_break(doc)
    
    # ==================== REAL-WORLD USE CASES ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '8. REAL-WORLD USE CASES', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    use_cases_intro = doc.add_paragraph()
    use_cases_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    use_cases_intro.add_run(
        'While the Snake Game is primarily an entertainment application, it has several practical '
        'applications and use cases in various domains:'
    )
    
    doc.add_paragraph()
    
    use_cases = [
        ('Educational Tool', 
         'The game serves as an excellent teaching resource for computer science students learning '
         'programming concepts such as loops, conditionals, data structures, object-oriented programming, '
         'and event handling. It provides a practical, engaging example of applying theoretical concepts.'),
        
        ('Game Development Learning', 
         'Aspiring game developers use Snake Game as a starting project to understand fundamental game '
         'development concepts including game loops, collision detection, state management, rendering, '
         'and user input handling before moving to more complex projects.'),
        
        ('Cognitive Training', 
         'The game can be used in cognitive training programs to improve hand-eye coordination, reaction '
         'time, spatial awareness, and strategic planning skills. It requires players to think ahead and '
         'plan their movements.'),
        
        ('Algorithm Demonstration', 
         'The Snake Game effectively demonstrates various algorithms including pathfinding (for AI players), '
         'collision detection, random number generation, and grid-based movement systems. It\'s often used '
         'in algorithm courses.'),
        
        ('User Interface Testing', 
         'The game can serve as a test application for evaluating input devices, display systems, and '
         'user interface responsiveness. It provides immediate visual feedback for input testing.'),
        
        ('Embedded Systems Projects', 
         'Simplified versions of Snake Game are commonly implemented on microcontrollers and embedded '
         'systems (Arduino, Raspberry Pi) to demonstrate graphics capabilities and input handling on '
         'constrained hardware.'),
        
        ('AI and Machine Learning', 
         'The game is frequently used as a test environment for reinforcement learning algorithms and '
         'AI agents. Researchers train AI models to play Snake Game autonomously, demonstrating machine '
         'learning concepts.'),
        
        ('Mobile App Development', 
         'The game logic can be adapted for mobile platforms, serving as a learning project for mobile '
         'app development, touch input handling, and cross-platform deployment.'),
        
        ('Recreational Gaming', 
         'As a casual game, it provides quick entertainment during breaks, helping with stress relief '
         'and mental relaxation. Its simple mechanics make it accessible to all age groups.'),
        
        ('Programming Competitions', 
         'Snake Game variations are used in coding competitions where participants implement AI bots '
         'to play the game optimally, competing for the highest scores or survival time.'),
    ]
    
    for title, description in use_cases:
        doc.add_heading(title, level=3)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(description)
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # ==================== FUTURE ENHANCEMENTS ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '9. FUTURE ENHANCEMENTS', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    enhancements_intro = doc.add_paragraph()
    enhancements_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    enhancements_intro.add_run(
        'The current implementation provides a solid foundation for the Snake Game. However, there are '
        'numerous opportunities for enhancement to improve gameplay, add features, and increase engagement:'
    )
    
    doc.add_paragraph()
    
    enhancements = [
        ('Difficulty Levels', [
            'Implement Easy, Medium, and Hard difficulty modes',
            'Adjust snake speed based on difficulty level',
            'Add obstacles in higher difficulty levels',
            'Implement progressive difficulty that increases with score',
        ]),
        
        ('Power-ups and Special Items', [
            'Speed boost power-up for temporary faster movement',
            'Slow-down power-up to reduce snake speed',
            'Score multiplier for bonus points',
            'Invincibility power-up for temporary collision immunity',
            'Shrink power-up to reduce snake length',
        ]),
        
        ('Visual Enhancements', [
            'Add animated sprites for snake and food',
            'Implement particle effects for food consumption',
            'Add background themes and customizable color schemes',
            'Include smooth movement animations between grid cells',
            'Add visual effects for power-ups and special events',
        ]),
        
        ('Audio Features', [
            'Background music during gameplay',
            'Sound effects for food consumption',
            'Game over sound effect',
            'Menu navigation sounds',
            'Volume control options',
        ]),
        
        ('Game Modes', [
            'Classic Mode: Current implementation',
            'Timed Mode: Achieve highest score within time limit',
            'Survival Mode: Avoid obstacles while collecting food',
            'Multiplayer Mode: Two players competing on same screen',
            'AI Challenge Mode: Compete against AI-controlled snake',
        ]),
        
        ('Scoring System', [
            'High score persistence using file storage or database',
            'Leaderboard showing top 10 scores',
            'Achievement system with unlockable badges',
            'Combo system for consecutive food collection',
            'Bonus points for speed and efficiency',
        ]),
        
        ('User Interface Improvements', [
            'Main menu with options for New Game, Settings, High Scores',
            'Settings menu for customization (controls, audio, graphics)',
            'In-game HUD showing additional statistics',
            'Tutorial mode for new players',
            'Customizable control schemes',
        ]),
        
        ('Technical Enhancements', [
            'Save and load game state functionality',
            'Replay system to watch previous games',
            'Performance optimization for larger grids',
            'Mobile touch control support',
            'Online multiplayer with networking',
        ]),
        
        ('Obstacles and Challenges', [
            'Static obstacles placed on the grid',
            'Moving obstacles that change position',
            'Walls that end the game on collision',
            'Portal tiles that teleport the snake',
            'Time-limited food items that disappear',
        ]),
        
        ('Analytics and Statistics', [
            'Track total games played',
            'Average score calculation',
            'Longest snake achieved',
            'Total playtime tracking',
            'Performance graphs and trends',
        ]),
    ]
    
    for category, items in enhancements:
        doc.add_heading(category, level=3)
        for item in items:
            para = doc.add_paragraph(item, style='List Bullet')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # ==================== FUTURE SCOPE ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '10. FUTURE SCOPE', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    scope_intro = doc.add_paragraph()
    scope_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    scope_intro.add_run(
        'The Snake Game project has significant potential for expansion and evolution. The future '
        'scope encompasses various directions for development, research, and commercial applications:'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.1 Platform Expansion', level=2)
    platform_para = doc.add_paragraph()
    platform_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    platform_para.add_run(
        'The game can be ported to multiple platforms including mobile devices (iOS and Android), '
        'web browsers using technologies like Pygame to JavaScript transpilers or HTML5 Canvas, '
        'gaming consoles, and smartwatches. Cross-platform development frameworks could enable '
        'simultaneous deployment across all platforms with shared codebase.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.2 Artificial Intelligence Integration', level=2)
    ai_para = doc.add_paragraph()
    ai_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ai_para.add_run(
        'Implementing AI agents using machine learning techniques such as Deep Q-Learning, Genetic '
        'Algorithms, or Neural Networks could create intelligent opponents or assistants. The game '
        'could serve as a research platform for studying reinforcement learning, pathfinding algorithms '
        '(A*, Dijkstra), and decision-making systems. AI could also provide adaptive difficulty that '
        'adjusts to player skill level.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.3 Social and Multiplayer Features', level=2)
    social_para = doc.add_paragraph()
    social_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    social_para.add_run(
        'Integration with social media platforms for score sharing, implementing real-time online '
        'multiplayer with matchmaking systems, creating tournaments and competitive leagues, adding '
        'friend systems and challenges, and developing cooperative game modes where players work '
        'together toward common goals.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.4 Educational Applications', level=2)
    edu_para = doc.add_paragraph()
    edu_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    edu_para.add_run(
        'The game can be transformed into an educational tool teaching programming concepts through '
        'visual scripting, mathematics through score calculations and geometry, logic through puzzle '
        'variations, and problem-solving skills. It could include level editors allowing students to '
        'create custom challenges, and integration with learning management systems for classroom use.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.5 Commercial Opportunities', level=2)
    commercial_para = doc.add_paragraph()
    commercial_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    commercial_para.add_run(
        'The game could be monetized through various models including free-to-play with advertisements, '
        'premium version with additional features, in-app purchases for cosmetic items and power-ups, '
        'subscription model for exclusive content, and licensing to educational institutions. Brand '
        'partnerships could introduce themed versions of the game.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.6 Accessibility Features', level=2)
    accessibility_para = doc.add_paragraph()
    accessibility_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    accessibility_para.add_run(
        'Future versions could include colorblind-friendly modes with alternative color schemes, '
        'screen reader support for visually impaired users, customizable control schemes for players '
        'with motor disabilities, adjustable game speed for different skill levels, and text-to-speech '
        'for game instructions and feedback.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.7 Data Analytics and Research', level=2)
    analytics_para = doc.add_paragraph()
    analytics_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    analytics_para.add_run(
        'Collecting anonymized gameplay data could enable research into player behavior patterns, '
        'difficulty balancing, user engagement metrics, and game design optimization. This data could '
        'inform academic research in human-computer interaction, game theory, and behavioral psychology.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('10.8 Virtual and Augmented Reality', level=2)
    vr_para = doc.add_paragraph()
    vr_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    vr_para.add_run(
        'The game concept could be adapted for VR headsets creating an immersive 3D snake experience, '
        'or AR applications where the game is played on real-world surfaces using smartphone cameras. '
        'This would represent a significant evolution of the classic 2D gameplay.'
    )
    
    add_page_break(doc)
    
    # ==================== SUMMARY ====================
    doc.add_paragraph('\n')
    add_heading_custom(doc, '11. SUMMARY', level=1, color=RGBColor(0, 0, 139))
    doc.add_paragraph('\n')
    
    summary_para1 = doc.add_paragraph()
    summary_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para1.add_run(
        'This project successfully demonstrates the implementation of a classic Snake Game using Python '
        'and Pygame library. The development process encompassed requirement analysis, system design, '
        'coding, testing, and documentation, providing a comprehensive learning experience in game '
        'development and software engineering principles.'
    )
    
    doc.add_paragraph()
    
    summary_para2 = doc.add_paragraph()
    summary_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para2.add_run(
        'The implemented game features smooth gameplay mechanics, intuitive controls, real-time collision '
        'detection, score tracking, and proper state management. The code is structured using object-oriented '
        'programming principles, making it maintainable, extensible, and easy to understand. The use of '
        'Python type hints and enumerations enhances code clarity and reduces potential errors.'
    )
    
    doc.add_paragraph()
    
    summary_para3 = doc.add_paragraph()
    summary_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para3.add_run(
        'Throughout the development process, several key concepts were applied including game loop architecture, '
        'event-driven programming, sprite rendering, collision detection algorithms, and user input handling. '
        'The project demonstrates how these fundamental concepts combine to create an engaging interactive '
        'application.'
    )
    
    doc.add_paragraph()
    
    summary_para4 = doc.add_paragraph()
    summary_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para4.add_run(
        'The game successfully meets all functional requirements including directional control, collision '
        'detection, food generation, score tracking, pause functionality, and restart capability. Non-functional '
        'requirements such as performance, usability, and reliability are also satisfied, with the game running '
        'smoothly at a consistent frame rate across different platforms.'
    )
    
    doc.add_paragraph()
    
    summary_para5 = doc.add_paragraph()
    summary_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para5.add_run(
        'The project has significant educational value, serving as an excellent learning resource for students '
        'studying programming, game development, and software engineering. It provides practical experience with '
        'Python programming, library integration, algorithm implementation, and problem-solving.'
    )
    
    doc.add_paragraph()
    
    summary_para6 = doc.add_paragraph()
    summary_para6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para6.add_run(
        'Future enhancements and scope for expansion are extensive, ranging from additional game features and '
        'visual improvements to platform expansion and AI integration. The modular code structure facilitates '
        'easy addition of new features without disrupting existing functionality.'
    )
    
    doc.add_paragraph()
    
    summary_para7 = doc.add_paragraph()
    summary_para7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para7.add_run(
        'In conclusion, this Snake Game project successfully achieves its objectives of creating a functional, '
        'entertaining game while demonstrating proficiency in Python programming and game development concepts. '
        'The project provides a solid foundation for further development and serves as a valuable portfolio piece '
        'showcasing technical skills and understanding of software development principles.'
    )
    
    doc.add_paragraph('\n' * 2)
    
    # Key Achievements
    doc.add_heading('Key Achievements', level=2)
    
    achievements = [
        'Successfully implemented a fully functional Snake Game with smooth gameplay',
        'Applied object-oriented programming principles for clean code architecture',
        'Implemented efficient collision detection and game state management',
        'Created an intuitive user interface with responsive controls',
        'Developed comprehensive documentation covering all project aspects',
        'Demonstrated proficiency in Python programming and Pygame library',
        'Designed a maintainable and extensible codebase for future enhancements',
        'Achieved cross-platform compatibility (Windows, macOS, Linux)',
    ]
    
    for achievement in achievements:
        para = doc.add_paragraph(achievement, style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph('\n' * 2)
    
    # Conclusion statement
    conclusion = doc.add_paragraph()
    conclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conclusion.add_run('* * * END OF DOCUMENTATION * * *').bold = True
    conclusion.runs[0].font.size = Pt(14)
    conclusion.runs[0].font.color.rgb = RGBColor(0, 0, 139)
    
    # Save the document
    doc.save('/Users/sheshavanand/MCA study materials/sem3/python/project/Snake_Game_Documentation.docx')
    print("Documentation created successfully: Snake_Game_Documentation.docx")

if __name__ == "__main__":
    create_documentation()
